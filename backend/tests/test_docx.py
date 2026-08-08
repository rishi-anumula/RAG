import io
import docx
import pytest
from app.services.docx_service import extract_text_from_docx

def test_extract_text_from_docx():
    """
    Tests text extraction from a sample docx document containing both
    paragraphs and table cells.
    """
    doc = docx.Document()
    doc.add_paragraph("First paragraph content.")
    doc.add_paragraph("Second paragraph content.")
    
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Cell 1"
    table.cell(0, 1).text = "Cell 2"
    table.cell(1, 0).text = "Cell 3"
    table.cell(1, 1).text = "Cell 4"

    buffer = io.BytesIO()
    doc.save(buffer)
    file_bytes = buffer.getvalue()

    extracted = extract_text_from_docx(file_bytes)

    assert "First paragraph content." in extracted
    assert "Second paragraph content." in extracted
    assert "Cell 1" in extracted
    assert "Cell 4" in extracted

    # Check joining with newlines
    lines = [line.strip() for line in extracted.split("\n") if line.strip()]
    assert len(lines) == 6
