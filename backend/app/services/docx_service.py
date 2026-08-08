import io
from typing import Union
import docx
from app.core.logging_config import get_logger

logger = get_logger(__name__)

def extract_text_from_docx(file_bytes: Union[bytes, str]) -> str:
    """
    Extracts text from a DOCX file provided either as raw bytes or a file path string.
    Opens docx.Document(io.BytesIO(file_bytes)) and loops through doc.paragraphs
    and doc.tables to collect all non-empty text, joining paragraphs with newlines.
    """
    try:
        if isinstance(file_bytes, (bytes, bytearray)):
            doc = docx.Document(io.BytesIO(file_bytes))
        elif isinstance(file_bytes, str):
            with open(file_bytes, "rb") as f:
                doc = docx.Document(io.BytesIO(f.read()))
        else:
            raise ValueError("Invalid file input type. Expected bytes or file path string.")

        extracted_parts = []

        # Extract text from document paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip() if paragraph.text else ""
            if text:
                extracted_parts.append(text)

        # Extract text from document tables without duplicate cells
        seen_cells = set()
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell._element in seen_cells:
                        continue
                    seen_cells.add(cell._element)
                    cell_text = cell.text.strip() if cell.text else ""
                    if cell_text:
                        extracted_parts.append(cell_text)

        full_text = "\n".join(extracted_parts)
        logger.info(f"Successfully extracted {len(full_text)} characters from DOCX document.")
        return full_text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX document: {e}", exc_info=True)
        raise ValueError(f"Could not parse DOCX file: {e}")

